import os
from google import genai
from docx import Document

class TranslationBackend:
    def __init__(self, api_key, translation_type='basic'):
        self.api_key = api_key
        self.translation_type = translation_type
        try:
            self.client = genai.Client(api_key=self.api_key)
        except Exception as e:
            print(f"Error configuring Gemini: {e}")
            self.client = None

    def process_docx(self, input_path, output_path, progress_callback=None):
        if not self.client:
            raise ValueError("Gemini client not initialized. Check API Key.")

        try:
            doc = Document(input_path)
            total_paragraphs = len(doc.paragraphs)
            
            for i, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue
                
                original_text = para.text
                # Optimization: Only send to Gemini if it contains parentheses with ascii chars inside
                if "(" in original_text and ")" in original_text:
                    processed_text = self.interact_with_gemini(original_text)
                    if processed_text and processed_text != original_text:
                        para.text = processed_text
                
                if progress_callback:
                    progress_callback(int((i / total_paragraphs) * 100))

            doc.save(output_path)
            # Finish progress
            if progress_callback:
                 progress_callback(100)
            return True, "Processing Complete"
        except Exception as e:
            return False, str(e)

    def interact_with_gemini(self, text):
        prompt = f"""
You are an expert linguist specializing in Urdu and English with an extremely sharp eye for etymology. Your task is to process the following text content while preserving all structure perfectly.

Your main goal is to meticulously identify *every single word* that originates from the English language and reformat it.
For this specific task, focus on **English words enclosed in parentheses**, e.g., `(Word)`.

1.  English words written in English script (e.g., "(Standard)").
2.  English words written in Urdu script, also known as loanwords (e.g., "سٹینڈرڈ"). You must be very aggressive in identifying these. Words like "کمپیوٹر", "فارمیٹنگ", "ڈاکومنٹ", "کاپی", "ایپلی کیشن", "پراسیس", etc., are all English loanwords and MUST be formatted. Do not treat them as pure Urdu.

Apply the following formatting rules with extreme precision to every identified English-origin term found in the context of parentheses:

**Rule 1: The Three-Part Format**
Every English term must be presented in this exact sequence:
[[Urdu Pronunciation]] (English Spelling) {{Urdu Translation}}

**Rule 2: Urdu Pronunciation [[...]]**
Write the word's pronunciation using Urdu script and enclose it in double square brackets. 
Check the ONE WORD immediately preceding the parentheses. 
- IF that preceding word is the Urdu pronunciation/transliteration of the English word, wrap it in `[[...]]`.
Example: If text has "سٹینڈرڈ (Standard)", output "[[سٹینڈرڈ]] (Standard)..."
Example: [[سٹینڈرڈ]], [[فارمیٹنگ]], [[کاپی]].

**Rule 3: English Spelling (...)**
Immediately follow the pronunciation with the original English spelling in parentheses. You must provide the correct English spelling.
Example: (Standard), (formatting), (copy).

**Rule 4: Contextual Urdu Translation {{...}}**
Finally, add the actual Urdu meaning in double curly brackets. 
- You MUST analyze the context to provide the most natural translation.
- If there are multiple appropriate Urdu words for the context, include them separated by a slash.
Example: {{معیار}}, {{ترتیب/ساخت}}, {{نقل/مصنفہ}}.

**Complete Transformation Example:**
The word "سٹینڈرڈ" in the text must become "[[سٹینڈرڈ]] (Standard) {{معیار}}"

Input: "یہ کنسٹرکشن (construction) کا کام ہے۔"
Output: "یہ [[کنسٹرکشن]] (construction) {{تعمیر}} کا کام ہے۔"

**Critical Constraints:**
-   Preserve all existing tags and structure perfectly.
-   Do not modify any text that is genuinely pure Urdu. Your focus is only on English and English-origin words.
-   Be thorough. It is a mistake to miss an English loanword like "سٹینڈرڈ".
-   Your final output MUST be only the processed text content. Do not include any markdown code fences or explanations.

Here is the content to process:
{text}
"""
        try:
            response = self.client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )
            result = response.text.strip()
            # Safety check: if result is empty, return original
            if not result:
                return text
            return result
        except Exception as e:
            print(f"Gemini API Error: {e}")
            return text 
