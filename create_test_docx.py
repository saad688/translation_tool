from docx import Document

doc = Document()
doc.add_paragraph("لیکن جب بندۂ مومن یہ ذمّہ داری لیتا ہے کہ اس کو خدا کے سامنے بہت ہی احتیاط کے ساتھ، صورتِ حال کے مطابق گفتگو کرنی ہے تو وہ ان لفظوں کے کنسٹرکشن (construction) میں اور اس احتیاط میں، اس اہتمام میں، اس کوشش میں مکمّل توجّہ کے ساتھ لگے رہے گا")
doc.add_paragraph("This is a simple test with a word like ہیلو (hello) world.")
doc.add_paragraph("Sometimes pronunciation matches like سکول (school) but sometimes it just translates like کتاب (book).")
doc.save("test_input.docx")
print("test_input.docx created")
